from docx import Document
import os
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_rows_by_t5(docx_file, entity_names):
    document = Document(docx_file)
    rows_dict = {entity_name: [] for entity_name in entity_names}
   
    prev_rows = [] 
    for table in document.tables:
        for row in table.rows:
            
            first_cell_value = row.cells[0].text.strip()
            # Check if the value matches any of the specified entity names
            


            for entity_name in entity_names:
                if first_cell_value.lower() == entity_name.lower():
                        if(entity_name=='Price'):
                                val =[cell.text.strip() for cell in prev_rows[0].cells]
                                rows_dict["Item No"].append(["Item No","",val[0][0:4]])
                                rows_dict["Item Description"].append(["Item Description","",val[0][4:]])
                                rows_dict[entity_name].append([cell.text.strip() for cell in row.cells])
                        elif(entity_name=="Quantity"):
                            val = [cell.text.strip() for cell in row.cells]
                            rows_dict["UOM"].append(["UOM","",val[2][2:]])
                            rows_dict["Quantity"].append(["Quantity","",val[2][0:2]])
                        else:
                            rows_dict[entity_name].append([cell.text.strip() for cell in row.cells])
                        break
            prev_rows.append(row)
            if len(prev_rows) > 2:
                prev_rows.pop(0) 
    ans =[]
    a=[]
    for x in entity_names:
        a.append(x)
    ans.append(a)
    for i in range(0,len(rows_dict["InternalNote"])+1):
        arr =[]
        for x in entity_names:
            if(i<len(rows_dict[x])):
                arr.append(rows_dict[x][i][2])
            else:
                arr.append('')
        ans.append(arr)
    return ans



def extract_rows_by_first_entity(docx_file, entity_names):
    document = Document(docx_file)
    rows_dict = {entity_name: [] for entity_name in entity_names}
    st =set()

    for table in document.tables:
        for row in table.rows:
            # Extract the value of the first cell in the row
            first_cell_value = row.cells[0].text.strip()
            i =0
            for x in first_cell_value:
                if(x==' '):
                    break
                i=i+1
            xx=first_cell_value[i:]
            if(xx.strip()=='Inco Term' or xx.strip()=='Inco Term Location'):
                first_cell_value=xx.strip()
                # rows_dict[first_cell_value].append([cell.text.strip() for cell in row.cells])

            # Check if the value matches any of the specified entity names
            for entity_name in entity_names:
                if first_cell_value.lower() == entity_name.lower() and entity_name not in st:
                    if(entity_name!="Quantity"):
                        rows_dict[entity_name].append([cell.text.strip() for cell in row.cells])
                    if(entity_name!="InternalNote" and entity_name != "Quantity"):
                        st.add(entity_name)
                    if(entity_name=="Quantity"):
                        rows_dict["InternalNote"].append([cell.text.strip() for cell in row.cells])
                    break
    return rows_dict

def write_to_new_docx(r1,r2,r3,r4,r2D, output_file):
    document = Document()
    t=[r1,r2,r3,r4]
    image_path = "t11.jpg"
    document.add_picture(image_path, width=Inches(2))
    heading = document.add_heading('CLIENT INQUIRY FORM', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    for i in range(0,4):
        for entity_name, rows in t[i].items():
            if rows:
                 table = document.add_table(rows=1, cols=len(rows[0]))
                 table.style = 'Table Grid'
                 for i, cell_value in enumerate(rows[0]):
                     table.cell(0, i).text = cell_value
                 for row in rows[1:]:
                     row_cells = table.add_row().cells
                     for i, cell_value in enumerate(row):
                         row_cells[i].text = cell_value
        document.add_paragraph()
    
    table = document.add_table(rows=len(r2D), cols=len(r2D[0]))
    table.style = 'Table Grid'
    widths = (Inches(1), Inches(2), Inches(1),Inches(1),Inches(1),Inches(1),Inches(2))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    

    for i in range(len(r2D)):
        for j in range(len(r2D[0])):
            table.cell(i, j).text = str(r2D[i][j])
    
    document.add_paragraph()

    if os.path.exists(output_file):
        os.remove(output_file)
    document.save(output_file)
    # copyFile("form1.docx",output_file)
    
    # doc_path = output_file
    # if os.path.exists(doc_path):
    # # Open the document with the default application
    #     os.system(f"start {doc_path}")
    # else:
    #     print("Error: File not found.")


def main():
    # Provide the path to your Word document
    docx_file = 'test1.docx'


    entity_names = ["Overview", "Owner", "Event Type", "Currency", "Timing Rules", "Publish time", "Due Date", "Currency Rules", "Allow Participants to select bidding currency","Inco Term", "Inco Term Location", "Requested Delivery Date","Quantity", "InternalNote" ]

    output_file = 'output.docx'  
    rows_dict = extract_rows_by_first_entity(docx_file, entity_names)
    write_to_new_docx(rows_dict, output_file)
    print("Rows extracted and written to new document successfully.")

if __name__ == "__main__":
    main()
